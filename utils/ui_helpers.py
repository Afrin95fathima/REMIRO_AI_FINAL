import streamlit as st
import base64
import re
import time
import random
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import io
import os

def convert_markdown_to_html(text):
    """Convert markdown formatting to clean HTML for display"""
    if not text:
        return ""
        
    # Convert headers with proper styling
    text = re.sub(r'^### (.*)', r'<h3 style="color: #5a6fd8; font-size: 18px; margin-top: 25px; margin-bottom: 10px;">\1</h3>', text, flags=re.MULTILINE)
    text = re.sub(r'^## (.*)', r'<h2 style="color: #4c63d2; font-size: 22px; margin-top: 30px; margin-bottom: 15px; border-left: 4px solid #667eea; padding-left: 15px;">\1</h2>', text, flags=re.MULTILINE)
    text = re.sub(r'^# (.*)', r'<h1 style="color: #667eea; font-size: 28px; text-align: center; margin-bottom: 30px; border-bottom: 3px solid #667eea; padding-bottom: 15px;">\1</h1>', text, flags=re.MULTILINE)
    
    # Convert bold text
    text = re.sub(r'\*\*(.*?)\*\*', r'<strong style="color: #4c63d2; font-weight: bold;">\1</strong>', text)
    
    # Convert italic text
    text = re.sub(r'\*(.*?)\*', r'<em style="color: #5a6fd8; font-style: italic;">\1</em>', text)
    
    # Convert bullet points with better styling
    text = re.sub(r'^- (.*)', r'<li style="margin: 8px 0; line-height: 1.6;">\1</li>', text, flags=re.MULTILINE)
    
    # Wrap consecutive list items in styled ul tags
    def replace_list(match):
        return f'<ul style="margin: 15px 0; padding-left: 25px;">\n{match.group(0)}</ul>\n'
    
    text = re.sub(r'(<li[^>]*>.*?</li>\s*)+', replace_list, text, flags=re.DOTALL)
    
    # Convert numbered lists
    text = re.sub(r'^(\d+)\. (.*)', r'<li style="margin: 8px 0; line-height: 1.6;">\2</li>', text, flags=re.MULTILINE)
    
    # Convert line breaks to proper spacing
    text = text.replace('\n\n', '</p><p style="margin: 15px 0; line-height: 1.8;">')
    text = text.replace('\n', '<br>')
    
    # Wrap content in paragraph tags
    if not text.startswith('<'):
        text = f'<p style="margin: 15px 0; line-height: 1.8;">{text}</p>'
    
    # Clean up extra br tags around headers and lists
    text = re.sub(r'<br>\s*(<h[1-3])', r'\1', text)
    text = re.sub(r'(<h[1-3][^>]*>.*?</h[1-3]>)\s*<br>', r'\1', text)
    text = re.sub(r'<br>\s*(<ul)', r'\1', text)
    text = re.sub(r'(</ul>)\s*<br>', r'\1', text)
    
    # Wrap everything in a styled container
    styled_content = f"""
    <div style="
        font-family: 'Times New Roman', Times, serif;
        line-height: 1.8;
        color: #333;
        background-color: #fdfdfd;
        padding: 30px;
        border-radius: 8px;
        border: 1px solid #f0f0f0;
        margin: 20px 0;
    ">
        {text}
    </div>
    """
    
    return styled_content

def render_interactive_question(question_data, key):
    """Render an interactive question with multiple choice options"""
    if not question_data or not isinstance(question_data, dict):
        return None
    
    question_text = question_data.get("question", "")
    question_type = question_data.get("type", "text")
    options = question_data.get("options", [])
    
    st.markdown(f"### 💭 {question_text}")
    
    if question_type == "multiple_choice" and options:
        st.markdown("**Select all that apply:**")
        
        selected_options = []
        
        # Create columns for better layout
        num_cols = min(2, len(options))  # Max 2 columns
        cols = st.columns(num_cols)
        
        for i, option in enumerate(options):
            col_index = i % num_cols
            with cols[col_index]:
                if st.checkbox(option, key=f"{key}_{i}"):
                    selected_options.append(option)
        
        # Add submit button
        if st.button("Submit Answers", key=f"{key}_submit", type="primary"):
            if selected_options:
                return selected_options
            else:
                st.warning("Please select at least one option.")
                return None
        
        return None  # Don't submit until button is pressed
    else:
        # Fallback to text input for other question types
        text_response = st.text_input("Your answer:", key=f"{key}_text")
        if st.button("Submit", key=f"{key}_submit", type="primary"):
            if text_response.strip():
                return text_response.strip()
            else:
                st.warning("Please provide an answer.")
                return None
        return None

def simulate_typing_delay(text_length):
    """Simulate natural typing delay based on text length"""
    # Base delay of 2-4 seconds + additional time based on text length
    base_delay = random.uniform(2.0, 4.0)
    length_delay = min(text_length / 50, 6.0)  # Max 6 seconds additional
    total_delay = base_delay + length_delay
    
    # Show typing indicator
    typing_placeholder = st.empty()
    typing_placeholder.markdown("🤔 *Thinking...*")
    time.sleep(total_delay)
    typing_placeholder.empty()

def render_multiple_choice_question(question, options, key):
    """Render a multiple choice question with checkboxes"""
    st.markdown(f"**{question}**")
    
    selected_options = []
    for i, option in enumerate(options):
        if st.checkbox(option, key=f"{key}_{i}"):
            selected_options.append(option)
    
    return selected_options

def format_final_report_display(text):
    """Specifically format final career reports for optimal display"""
    if not text:
        return "No report content available."
    
    # Ensure text is a string
    if isinstance(text, tuple):
        text = str(text[0]) if text else ""
    elif not isinstance(text, str):
        text = str(text)
    
    # Convert markdown to styled HTML with enhanced styling for reports
    formatted_html = convert_markdown_to_html(text)
    
    # Add a special wrapper for final reports
    final_report_html = f"""
    <div style="
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2px;
        border-radius: 12px;
        margin: 20px 0;
    ">
        <div style="
            background: white;
            border-radius: 10px;
            overflow: hidden;
        ">
            <div style="
                background: linear-gradient(90deg, #667eea, #764ba2);
                color: white;
                text-align: center;
                padding: 15px;
                margin-bottom: 0;
            ">
                <h2 style="
                    margin: 0;
                    font-family: 'Times New Roman', serif;
                    font-size: 24px;
                    font-weight: bold;
                ">🚀 Your Career Strategy Report</h2>
            </div>
            {formatted_html}
        </div>
    </div>
    """
    
    return final_report_html

def format_report_for_display(text):
    """Format report text for better display in Streamlit"""
    if not text:
        return ""
        
    # Convert to HTML format for better display
    return convert_markdown_to_html(text)

def generate_pdf_report(report_text, user_details):
    """Generate PDF report"""
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=1*Inches, bottomMargin=1*Inches)
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            textColor=HexColor('#667eea'),
            alignment=1  # Center alignment
        )
        
        header_style = ParagraphStyle(
            'CustomHeader',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            spaceBefore=20,
            textColor=HexColor('#4c63d2'),
            leftIndent=0
        )
        
        subheader_style = ParagraphStyle(
            'CustomSubHeader',
            parent=styles['Heading3'],
            fontSize=14,
            spaceAfter=10,
            spaceBefore=15,
            textColor=HexColor('#5a6fd8')
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=6,
            fontName='Times-Roman'
        )
        
        # Build the story
        story = []
        
        # Title
        story.append(Paragraph("🚀 Your Personalized Career Strategy Blueprint", title_style))
        story.append(Spacer(1, 12))
        
        # User info
        story.append(Paragraph(f"<b>Prepared for:</b> {user_details.get('name', 'User')}", normal_style))
        story.append(Paragraph(f"<b>Generated by:</b> Remiro AI Career Guidance Assistant", normal_style))
        story.append(Spacer(1, 20))
        
        # Process the report content
        lines = report_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
                continue
                
            if line.startswith('### '):
                story.append(Paragraph(line[4:], subheader_style))
            elif line.startswith('## '):
                story.append(Paragraph(line[3:], header_style))
            elif line.startswith('# '):
                story.append(Paragraph(line[2:], title_style))
            elif line.startswith('- '):
                story.append(Paragraph(f"• {line[2:]}", normal_style))
            else:
                # Clean up any remaining markdown
                clean_line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
                clean_line = re.sub(r'\*(.*?)\*', r'<i>\1</i>', clean_line)
                story.append(Paragraph(clean_line, normal_style))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Error generating PDF: {str(e)}")
        return None

def generate_docx_report(report_text, user_details):
    """Generate DOCX report"""
    try:
        doc = Document()
        
        # Add title
        title = doc.add_heading('🚀 Your Personalized Career Strategy Blueprint', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add user info
        doc.add_paragraph()
        user_info = doc.add_paragraph()
        user_info.add_run(f"Prepared for: ").bold = True
        user_info.add_run(user_details.get('name', 'User'))
        
        gen_info = doc.add_paragraph()
        gen_info.add_run("Generated by: ").bold = True
        gen_info.add_run("Remiro AI Career Guidance Assistant")
        
        doc.add_paragraph()
        
        # Process the report content
        lines = report_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
                
            if line.startswith('### '):
                heading = doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                heading = doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                heading = doc.add_heading(line[2:], level=1)
            elif line.startswith('- '):
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                # Clean up markdown formatting
                clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', line[2:])
                clean_text = re.sub(r'\*(.*?)\*', r'\1', clean_text)
                p.add_run(clean_text)
            else:
                p = doc.add_paragraph()
                # Clean up markdown formatting and add styled text
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        run = p.add_run(part[1:-1])
                        run.italic = True
                    else:
                        p.add_run(part)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Error generating DOCX: {str(e)}")
        return None
import json

def setup_page_config(title):
    """Setup page configuration"""
    st.set_page_config(
        page_title=title,
        page_icon="🧭",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Apply simple, clean CSS with Times New Roman font
    st.markdown("""
        <style>
            /* Main app styling */
            .stApp {
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
                font-family: "Times New Roman", Times, serif !important;
            }
            
            /* Main content area - simplified */
            .main .block-container {
                background-color: #ffffff !important;
                border-radius: 15px !important;
                padding: 2rem !important;
                margin: 1rem !important;
                box-shadow: 0 10px 30px rgba(0, 0, 0, 0.1) !important;
            }
            
            /* Headers */
            h1, h2, h3, h4, h5, h6 {
                font-family: "Times New Roman", Times, serif !important;
                color: #2d3748 !important;
                font-weight: 600 !important;
                margin: 1rem 0 !important;
            }
            
            h1 {
                text-align: center !important;
                color: #667eea !important;
                font-size: 2.5rem !important;
            }
            
            /* All text elements */
            .stApp *, p, div, span, label, input, select, textarea {
                font-family: "Times New Roman", Times, serif !important;
                color: #2d3748 !important;
            }
            
            /* Form elements - simple and clean */
            .stTextInput input, .stSelectbox select, .stNumberInput input, .stTextArea textarea {
                background-color: #ffffff !important;
                border: 2px solid #e2e8f0 !important;
                border-radius: 8px !important;
                padding: 12px !important;
                color: #2d3748 !important;
                font-size: 16px !important;
                font-family: "Times New Roman", Times, serif !important;
            }
            
            .stTextInput input:focus, .stSelectbox select:focus, .stNumberInput input:focus, .stTextArea textarea:focus {
                border-color: #667eea !important;
                box-shadow: 0 0 0 2px rgba(102, 126, 234, 0.2) !important;
                outline: none !important;
            }
            
            /* Buttons */
            .stButton button {
                background-color: #667eea !important;
                color: white !important;
                border: none !important;
                border-radius: 8px !important;
                padding: 12px 24px !important;
                font-family: "Times New Roman", Times, serif !important;
                font-weight: 600 !important;
                font-size: 16px !important;
            }
            
            .stButton button:hover {
                background-color: #5a67d8 !important;
            }
            
            /* Slider styling */
            .stSlider > div {
                padding: 1rem 0 !important;
            }
            
            .stSlider [data-testid="stThumbValue"] {
                background-color: #667eea !important;
                border-color: #667eea !important;
                color: white !important;
                font-family: "Times New Roman", Times, serif !important;
            }
            
            /* Chat messages */
            .stChatMessage {
                background-color: #f7fafc !important;
                border-radius: 10px !important;
                padding: 1rem !important;
                margin: 1rem 0 !important;
                border-left: 4px solid #667eea !important;
            }
            
            /* Sidebar */
            .css-1d391kg {
                background: linear-gradient(135deg, #667eea, #764ba2) !important;
                color: white !important;
                padding: 1rem !important;
            }
            
            .css-1d391kg h1, .css-1d391kg h2, .css-1d391kg h3 {
                color: white !important;
                font-family: "Times New Roman", Times, serif !important;
            }
            
            .css-1d391kg p, .css-1d391kg div {
                color: white !important;
                font-family: "Times New Roman", Times, serif !important;
            }
            
            /* Remove extra styling for simplicity */
            .stMarkdown {
                font-family: "Times New Roman", Times, serif !important;
            }
            
            /* Ensure all elements use Times New Roman */
            * {
                font-family: "Times New Roman", Times, serif !important;
            }
        </style>
    """, unsafe_allow_html=True)

def render_moving_background():
    """Simple background - no complex animations"""
    pass

def render_chat_interface(master_agent, user_details):
    """Render the enhanced chat interface with interactive questions"""
    
    # Initialize current agent state if not present
    if "current_agent_name" not in st.session_state:
        current_agent = master_agent.get_current_agent()
        if current_agent:
            st.session_state.current_agent_name = master_agent.agent_sequence[master_agent.current_agent_index]
        else:
            st.session_state.current_agent_name = "master"
    
    # Initialize question state
    if "current_question" not in st.session_state:
        st.session_state.current_question = None
        st.session_state.current_question_index = 0
        st.session_state.waiting_for_answer = False
    
    # Show current agent indicator
    current_agent_name = st.session_state.current_agent_name.replace("_", " ").title()
    
    if master_agent.final_integration_started:
        st.markdown("### 🧠 **Creating Your Career Strategy**")
    else:
        st.markdown(f"### 👤 **Current Focus:** {current_agent_name} Analysis")
    
    # Display chat messages from history
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            if message["role"] == "assistant":
                # Check if this message contains a question that needs interactive handling
                if "current_question_data" in message and message.get("interactive", False):
                    # This is an interactive question - render it specially
                    question_data = message["current_question_data"]
                    st.markdown(message["content"])
                else:
                    # Regular message - display normally
                    st.markdown(message["content"])
            else:
                st.markdown(message["content"])
    
    # Handle interactive question flow
    if st.session_state.waiting_for_answer and st.session_state.current_question:
        question_data = st.session_state.current_question
        question_key = f"q_{st.session_state.current_question_index}_{time.time()}"
        
        # Render the interactive question
        user_response = render_interactive_question(question_data, question_key)
        
        if user_response is not None:
            # User submitted an answer
            st.session_state.waiting_for_answer = False
            
            # Process the response
            if isinstance(user_response, list):
                response_text = ", ".join(user_response)
            else:
                response_text = str(user_response)
            
            # Add user message to chat history
            st.session_state.messages.append({"role": "user", "content": f"Selected: {response_text}"})
            
            # Process with master agent
            with st.chat_message("assistant"):
                typing_placeholder = st.empty()
                typing_placeholder.markdown("🤔 *Analyzing your response...*")
                
                try:
                    # Process the message through the master agent
                    response_data = master_agent.process_message(response_text, user_details)
                    
                    # Simulate natural thinking time
                    thinking_time = random.uniform(3.0, 6.0)
                    time.sleep(thinking_time)
                    
                    typing_placeholder.empty()
                    
                    # Handle response
                    if isinstance(response_data, tuple):
                        if len(response_data) == 2:
                            response, is_complete = response_data
                            results = None
                        elif len(response_data) == 3:
                            response, is_complete, results = response_data
                        else:
                            response = str(response_data)
                            is_complete = False
                            results = None
                    else:
                        response = str(response_data)
                        is_complete = False
                        results = None
                    
                    # Display the response
                    st.markdown(response)
                    
                    # Add response to chat history
                    st.session_state.messages.append({"role": "assistant", "content": response})
                    
                    # Check if we need to ask another question
                    current_agent = master_agent.get_current_agent()
                    if current_agent and not is_complete:
                        # Get next question from current agent
                        next_question = get_next_question_from_agent(current_agent, st.session_state.current_question_index + 1)
                        if next_question:
                            st.session_state.current_question = next_question
                            st.session_state.current_question_index += 1
                            st.session_state.waiting_for_answer = True
                            st.rerun()
                    
                    # Handle agent completion
                    if is_complete:
                        handle_agent_completion(master_agent, results)
                        
                except Exception as e:
                    typing_placeholder.empty()
                    st.error(f"An error occurred: {str(e)}")
            
            st.rerun()
    
    # Regular chat input for non-interactive questions or initial messages
    elif not st.session_state.waiting_for_answer:
        if prompt := st.chat_input("Type your response here..."):
            # Add user message to chat history
            st.session_state.messages.append({"role": "user", "content": prompt})
            
            # Display user message
            with st.chat_message("user"):
                st.markdown(prompt)
            
            # Process the message
            process_regular_message(master_agent, user_details, prompt)

def get_next_question_from_agent(agent, question_index):
    """Get the next question from an agent"""
    if hasattr(agent, 'questions') and question_index < len(agent.questions):
        return agent.questions[question_index]
    return None

def handle_agent_completion(master_agent, results):
    """Handle when an agent completes its analysis"""
    if results and not master_agent.final_integration_started:
        # Move to next agent
        next_agent = master_agent.move_to_next_agent()
        
        if next_agent:
            # Update current agent name
            st.session_state.current_agent_name = master_agent.agent_sequence[master_agent.current_agent_index]
            
            # Start with first question of next agent
            first_question = get_next_question_from_agent(next_agent, 0)
            if first_question:
                st.session_state.current_question = first_question
                st.session_state.current_question_index = 0
                st.session_state.waiting_for_answer = True
            
            # Add transition message
            transition_msg = f"Great! Now let's move to the **{st.session_state.current_agent_name.replace('_', ' ').title()}** analysis."
            st.session_state.messages.append({"role": "assistant", "content": transition_msg})
        else:
            # We've gone through all agents, moving to final integration
            st.session_state.current_agent_name = "master"
            st.session_state.current_question = None
            st.session_state.waiting_for_answer = False
    
    # If this is the final response, display download button
    if master_agent.final_report_generated:
        # Get the last assistant message as the report
        last_message = None
        for msg in reversed(st.session_state.messages):
            if msg["role"] == "assistant":
                last_message = msg["content"]
                break
        
        if last_message:
            offer_report_download(last_message, st.session_state.get("user_details", {}))

def process_regular_message(master_agent, user_details, prompt):
    """Process a regular text message"""
    with st.chat_message("assistant"):
        typing_placeholder = st.empty()
        typing_placeholder.markdown("🤔 *Analyzing your response...*")
        
        try:
            # Process the message through the master agent
            response_data = master_agent.process_message(prompt, user_details)
            
            # Simulate natural thinking time
            thinking_time = random.uniform(5.0, 10.0)
            time.sleep(thinking_time)
            
            typing_placeholder.empty()
            
            # Handle response structure
            if isinstance(response_data, tuple):
                if len(response_data) == 2:
                    response, is_complete = response_data
                    results = None
                elif len(response_data) == 3:
                    response, is_complete, results = response_data
                else:
                    response = str(response_data)
                    is_complete = False
                    results = None
            else:
                response = str(response_data)
                is_complete = False
                results = None
            
            # Display the response with proper formatting
            if master_agent.final_report_generated:
                # This is the final report - display as formatted HTML
                formatted_html = format_final_report_display(response)
                st.markdown(formatted_html, unsafe_allow_html=True)
            else:
                # Check if response contains a question that should be interactive
                current_agent = master_agent.get_current_agent()
                if current_agent and hasattr(current_agent, 'questions') and not is_complete:
                    # This might be a question - check if we should make it interactive
                    next_question = get_next_question_from_agent(current_agent, 0)
                    if next_question:
                        st.session_state.current_question = next_question
                        st.session_state.current_question_index = 0
                        st.session_state.waiting_for_answer = True
                        
                        # Display the question text but mark it for interactive handling
                        st.markdown(response)
                        st.session_state.messages.append({
                            "role": "assistant", 
                            "content": response,
                            "interactive": True,
                            "current_question_data": next_question
                        })
                        st.rerun()
                        return
                
                # Regular response
                st.markdown(response)
            
            # Add response to chat history
            st.session_state.messages.append({"role": "assistant", "content": response})
            
            # Handle completion
            if is_complete:
                handle_agent_completion(master_agent, results)
                
        except Exception as e:
            typing_placeholder.empty()
            st.error(f"An error occurred: {str(e)}")
            st.session_state.messages.append({
                "role": "assistant", 
                "content": "I apologize for the error. Please try rephrasing your response."
            })

def offer_report_download(report_text, user_details):
    """Offer a download button for the career report"""
    
    # Ensure report_text is a string
    if isinstance(report_text, tuple):
        report_text = str(report_text[0]) if report_text else "No report generated"
    elif not isinstance(report_text, str):
        report_text = str(report_text) if report_text else "No report generated"
    
    # Convert markdown to clean HTML
    formatted_content = convert_markdown_to_html(report_text)
    
    report_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Career Guidance Report</title>
        <style>
            body {{
                font-family: "Times New Roman", Times, serif;
                line-height: 1.8;
                color: #333;
                max-width: 900px;
                margin: 0 auto;
                padding: 40px;
                background-color: #fff;
            }}
            
            h1 {{
                color: #667eea;
                font-size: 28px;
                text-align: center;
                margin-bottom: 30px;
                border-bottom: 3px solid #667eea;
                padding-bottom: 15px;
            }}
            
            h2 {{
                color: #4c63d2;
                font-size: 22px;
                margin-top: 30px;
                margin-bottom: 15px;
                border-left: 4px solid #667eea;
                padding-left: 15px;
            }}
            
            h3 {{
                color: #5a6fd8;
                font-size: 18px;
                margin-top: 25px;
                margin-bottom: 10px;
            }}
            
            ul {{
                margin: 15px 0;
                padding-left: 25px;
            }}
            
            li {{
                margin: 8px 0;
                line-height: 1.6;
            }}
            
            .header-info {{
                text-align: center;
                background-color: #f8f9ff;
                padding: 20px;
                border-radius: 10px;
                margin-bottom: 30px;
                border: 1px solid #e0e6ff;
            }}
            
            .footer {{
                text-align: center;
                margin-top: 40px;
                padding-top: 20px;
                border-top: 2px solid #e0e6ff;
                font-style: italic;
                color: #666;
            }}
            
            strong {{
                color: #4c63d2;
                font-weight: bold;
            }}
            
            em {{
                color: #5a6fd8;
                font-style: italic;
            }}
            
            .content {{
                background-color: #fdfdfd;
                padding: 30px;
                border-radius: 8px;
                border: 1px solid #f0f0f0;
            }}
        </style>
    </head>
    <body>
        <div class="header-info">
            <h1>🚀 Your Personalized Career Strategy Blueprint</h1>
            <p><strong>Prepared for:</strong> {user_details.get('name', 'User')}</p>
            <p><strong>Date:</strong> {st.session_state.get('current_date', 'Today')}</p>
        </div>
        
        <div class="content">
            {formatted_content}
        </div>
        
        <div class="footer">
            <p><em>Generated by Remiro AI Career Guidance Assistant</em></p>
            <p><em>Your AI-Powered Career Companion</em></p>
        </div>
    </body>
    </html>
    """
    
    # Create download buttons for multiple formats
    st.markdown("---")
    st.markdown("### 📥 Download Your Career Report")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        # HTML Download
        b64_html = base64.b64encode(report_html.encode()).decode()
        href_html = f'<a href="data:text/html;base64,{b64_html}" download="career_report.html" style="text-decoration: none;"><div style="background-color: #667eea; color: white; padding: 10px 20px; border-radius: 5px; text-align: center; margin: 5px;">📄 HTML Report</div></a>'
        st.markdown(href_html, unsafe_allow_html=True)
    
    with col2:
        # PDF Download
        pdf_data = generate_pdf_report(report_text, user_details)
        if pdf_data:
            b64_pdf = base64.b64encode(pdf_data).decode()
            href_pdf = f'<a href="data:application/pdf;base64,{b64_pdf}" download="career_report.pdf" style="text-decoration: none;"><div style="background-color: #dc3545; color: white; padding: 10px 20px; border-radius: 5px; text-align: center; margin: 5px;">� PDF Report</div></a>'
            st.markdown(href_pdf, unsafe_allow_html=True)
        else:
            st.error("PDF generation failed")
    
    with col3:
        # DOCX Download
        docx_data = generate_docx_report(report_text, user_details)
        if docx_data:
            b64_docx = base64.b64encode(docx_data).decode()
            href_docx = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_docx}" download="career_report.docx" style="text-decoration: none;"><div style="background-color: #28a745; color: white; padding: 10px 20px; border-radius: 5px; text-align: center; margin: 5px;">📝 Word Document</div></a>'
            st.markdown(href_docx, unsafe_allow_html=True)
        else:
            st.error("DOCX generation failed")
    
    st.markdown("Choose your preferred format to download your personalized career report!")

def display_agent_complete_badge(agent_name):
    """Display completion badge for an agent"""
    st.markdown(f"✅ **{agent_name} Analysis Complete**")